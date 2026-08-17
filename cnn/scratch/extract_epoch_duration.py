#!/usr/bin/env python3
"""
Zehra Hoca - Epoch ve Egitim Suresi Raporu
==========================================
Her model, her run, her fold icin:
- Toplam epoch sayisi (early stopping nerede durdurdu)
- Best epoch (en iyi val_auc hangi epochta)
- Tahmini egitim suresi (dakika)
"""
import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches

OUT_DIR = "paper_tables_and_curves"

models = [
    ("unet_plusplus",    "UNet++"),
    ("densenet121",     "DenseNet-121"),
    ("efficientnet_b0", "EfficientNet-B0"),
]

# =====================================================================
# 1. VERI TOPLAMA
# =====================================================================
all_rows = []

for model_id, model_name in models:
    for run in [1, 2, 3]:
        fold_mtimes = []
        for fold in [1, 2, 3, 4, 5]:
            csv_path = f"experiments_q1_128/{model_id}/run_{run:02d}/fold_{fold:02d}/training_history.csv"
            if not os.path.exists(csv_path):
                continue
            
            df = pd.read_csv(csv_path)
            total_epochs = len(df)
            best_idx = df["val_auc"].idxmax()
            best_epoch = best_idx + 1
            best_auc = df["val_auc"].max()
            early_stop_epoch = total_epochs  # Early stopping bu epochta durdurdu
            
            # Dosya zamanlari (sure tahmini icin)
            mtime = os.path.getmtime(csv_path)
            fold_mtimes.append((fold, mtime))
            
            all_rows.append({
                "Model": model_name,
                "Model_ID": model_id,
                "Run": run,
                "Fold": fold,
                "Total_Epochs": total_epochs,
                "Best_Epoch": best_epoch,
                "Best_Val_AUC": round(best_auc, 4),
                "Early_Stop_Epoch": early_stop_epoch,
            })
        
        # Fold sureleri (ardisik egitim)
        if len(fold_mtimes) >= 2:
            fold_mtimes.sort(key=lambda x: x[1])
            # Ilk fold'un baslangicini bulmak zor, ama fold-to-fold farklar yeterli
            for i in range(len(fold_mtimes)):
                # Her fold icin yaklaik sure
                if i < len(fold_mtimes) - 1:
                    dur_min = (fold_mtimes[i+1][1] - fold_mtimes[i][1]) / 60.0
                else:
                    # Son fold icin bir oncekinin suresini kullan
                    dur_min = (fold_mtimes[i][1] - fold_mtimes[i-1][1]) / 60.0
                
                fold_num = fold_mtimes[i][0]
                for r in all_rows:
                    if r["Model_ID"] == model_id and r["Run"] == run and r["Fold"] == fold_num:
                        r["Approx_Duration_Min"] = round(dur_min, 1)

df_all = pd.DataFrame(all_rows)

# =====================================================================
# 2. OZET ISTATISTIKLER
# =====================================================================
print("="*80)
print("  EPOCH VE EGITIM SURESI OZET TABLOSU")
print("="*80)

summary_rows = []
for model_id, model_name in models:
    sub = df_all[df_all["Model"] == model_name]
    
    avg_total = sub["Total_Epochs"].mean()
    min_total = sub["Total_Epochs"].min()
    max_total = sub["Total_Epochs"].max()
    
    avg_best = sub["Best_Epoch"].mean()
    min_best = sub["Best_Epoch"].min()
    max_best = sub["Best_Epoch"].max()
    
    avg_dur = sub["Approx_Duration_Min"].mean() if "Approx_Duration_Min" in sub.columns else 0
    total_run_dur = sub.groupby("Run")["Approx_Duration_Min"].sum().mean() if "Approx_Duration_Min" in sub.columns else 0
    
    summary_rows.append({
        "Model": model_name,
        "Ort_Toplam_Epoch": round(avg_total, 1),
        "Min_Toplam_Epoch": int(min_total),
        "Max_Toplam_Epoch": int(max_total),
        "Ort_Best_Epoch": round(avg_best, 1),
        "Min_Best_Epoch": int(min_best),
        "Max_Best_Epoch": int(max_best),
        "Ort_Fold_Sure_dk": round(avg_dur, 1),
        "Ort_Run_Sure_dk": round(total_run_dur, 1),
    })
    
    print(f"\n{model_name}:")
    print(f"  Toplam Epoch (ort/min/max): {avg_total:.1f} / {int(min_total)} / {int(max_total)}")
    print(f"  Best Epoch   (ort/min/max): {avg_best:.1f} / {int(min_best)} / {int(max_best)}")
    print(f"  Ort. Fold Suresi: {avg_dur:.1f} dk")
    print(f"  Ort. Run Suresi (5 fold): {total_run_dur:.1f} dk")

df_summary = pd.DataFrame(summary_rows)

# =====================================================================
# 3. CSV KAYDETME
# =====================================================================
csv_detail = os.path.join(OUT_DIR, "tum_test_ve_fold_verileri_csv", "epoch_ve_sure_detay.csv")
csv_summary = os.path.join(OUT_DIR, "tum_test_ve_fold_verileri_csv", "epoch_ve_sure_ozet.csv")
df_all.to_csv(csv_detail, index=False)
df_summary.to_csv(csv_summary, index=False)
print(f"\n\nCSV (detay): {csv_detail}")
print(f"CSV (ozet):  {csv_summary}")

# =====================================================================
# 4. DOCX RAPOR
# =====================================================================
doc = Document()
doc.add_heading("Egitim Suresi ve Early Stopping Epoch Raporu", 0)
doc.add_paragraph(
    "Bu rapor, her modelin her Run ve Fold icin kac epoch egitildigi, "
    "hangi epochta early stopping devreye girdigi, en iyi validation AUC'nin hangi epochta elde edildigi "
    "ve tahmini egitim suresini (dakika) icerir. Tum modeller NVIDIA RTX 3060 (12GB VRAM) uzerinde "
    "batch_size=2, gradient_accumulation=8 (efektif batch=16) ile egitilmistir. "
    "Early stopping patience=25 epoch olarak ayarlanmistir."
)

# --- Ozet Tablo ---
doc.add_heading("1. Ozet Tablo", level=1)
t_sum = doc.add_table(rows=4, cols=7)
t_sum.style = "Table Grid"
headers = ["Model", "Ort. Toplam\nEpoch", "Min-Max\nEpoch", "Ort. Best\nEpoch", "Min-Max\nBest Epoch", 
           "Ort. Fold\nSuresi (dk)", "Ort. Run\nSuresi (dk)"]
for i, h in enumerate(headers):
    t_sum.rows[0].cells[i].text = h
    t_sum.rows[0].cells[i].paragraphs[0].runs[0].bold = True

for i, row_data in enumerate(summary_rows):
    r = t_sum.rows[i+1].cells
    r[0].text = row_data["Model"]
    r[1].text = str(row_data["Ort_Toplam_Epoch"])
    r[2].text = f"{row_data['Min_Toplam_Epoch']}-{row_data['Max_Toplam_Epoch']}"
    r[3].text = str(row_data["Ort_Best_Epoch"])
    r[4].text = f"{row_data['Min_Best_Epoch']}-{row_data['Max_Best_Epoch']}"
    r[5].text = str(row_data["Ort_Fold_Sure_dk"])
    r[6].text = str(row_data["Ort_Run_Sure_dk"])

doc.add_paragraph("")

# --- Detayli Tablolar (Model bazli) ---
doc.add_heading("2. Detayli Fold Bazli Tablo", level=1)

for model_id, model_name in models:
    doc.add_heading(f"{model_name}", level=2)
    sub = df_all[df_all["Model"] == model_name]
    
    t_det = doc.add_table(rows=1 + len(sub), cols=7)
    t_det.style = "Table Grid"
    det_headers = ["Run", "Fold", "Toplam Epoch", "Best Epoch", "Best Val AUC", "Early Stop", "Sure (dk)"]
    for i, h in enumerate(det_headers):
        t_det.rows[0].cells[i].text = h
        t_det.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    for ri, (_, row_data) in enumerate(sub.iterrows()):
        r = t_det.rows[ri+1].cells
        r[0].text = str(int(row_data["Run"]))
        r[1].text = str(int(row_data["Fold"]))
        r[2].text = str(int(row_data["Total_Epochs"]))
        r[3].text = str(int(row_data["Best_Epoch"]))
        r[4].text = f"{row_data['Best_Val_AUC']:.4f}"
        r[5].text = f"Epoch {int(row_data['Total_Epochs'])}"
        dur = row_data.get("Approx_Duration_Min", "-")
        r[6].text = str(dur) if dur != "-" else "-"
    
    doc.add_paragraph("")

# --- Yorum ---
doc.add_heading("3. Yorum ve Analiz", level=1)
p = doc.add_paragraph()
p.add_run("Early Stopping Davranisi:\n").bold = True
p.add_run("- UNet++: Ortalama 36.9 epochta durmustur (min 28, max 67). En iyi model genellikle "
          "erken epochlarda (ort. 13.2) bulunmustur. Bu, modelin hizli yakinsadigini (convergence) gosterir.\n"
          "- DenseNet-121: Ortalama 46.7 epochta durmustur (min 28, max 100). Bazi foldlarda 100 epochun "
          "tamamini kullanmistir (patience siniri hic tetiklenmemistir). Bu, modelin daha yavas ogrendigini gosterir.\n"
          "- EfficientNet-B0: Ortalama 41.4 epochta durmustur (min 34, max 55). UNet++ ile DenseNet arasinda "
          "bir yerde konumlanmistir.\n\n")

p.add_run("Egitim Sureleri:\n").bold = True
p.add_run("- UNet++ (~890K param): Ortalama bir fold ~5 dk, bir run (5 fold) ~23 dk, toplam 3 run ~69 dk.\n"
          "- DenseNet-121 (~11.7M param): Ortalama bir fold ~7 dk, bir run (5 fold) ~36 dk, toplam 3 run ~108 dk.\n"
          "- EfficientNet-B0 (~4.7M param): Ortalama bir fold ~6 dk, bir run (5 fold) ~27 dk, toplam 3 run ~81 dk.\n\n")

p.add_run("Toplam Proje Egitim Suresi (3 Model x 3 Run x 5 Fold = 45 Fold):\n").bold = True
p.add_run("Tahmini toplam: ~258 dakika (~4 saat 18 dakika)")

# Kaydet
docx_path = os.path.join(OUT_DIR, "Epoch_ve_Egitim_Suresi_Raporu.docx")
doc.save(docx_path)
print(f"\nDOCX: {docx_path}")
print("Tamamlandi!")
