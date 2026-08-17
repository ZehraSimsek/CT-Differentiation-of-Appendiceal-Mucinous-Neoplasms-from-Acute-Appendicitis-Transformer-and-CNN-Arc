#!/usr/bin/env python3
"""
Kapsamli Model Analiz Raporu Olusturucu
- Augmentasyonlu ve Augmentasyonsuz hiperparametre karsilastirmasi
- Grafik uretimi (3 adet)
- Detayli model bazli analiz
"""

import os
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT_DIR = "paper_tables_and_curves"
os.makedirs(OUT_DIR, exist_ok=True)

# =====================================================================
# DATA
# =====================================================================
models = ["UNet++", "DenseNet-121", "EfficientNet-B0"]

# --- Augmentasyonsuz (Eski) Sonuclar (Dis Test Ensemble) ---
aug_yok = {
    "UNet++":          {"AUC": 0.863, "Sens": 0.789, "Spec": 0.800, "Acc": 0.792, "F1": 0.857},
    "DenseNet-121":    {"AUC": 0.842, "Sens": 0.842, "Spec": 0.600, "Acc": 0.792, "F1": 0.865},
    "EfficientNet-B0": {"AUC": 0.842, "Sens": 1.000, "Spec": 0.000, "Acc": 0.792, "F1": 0.884},
}

# --- Augmentasyonlu (Nihai 3-Run Ortalamalari) Sonuclar ---
# Run ortalamalari hesaplanmis metrikler (master CSV'den)
aug_var = {
    "UNet++":          {"AUC": 0.853, "Sens": 0.596, "Spec": 0.933, "Acc": 0.680, "F1": 0.747},
    "DenseNet-121":    {"AUC": 0.849, "Sens": 0.596, "Spec": 0.933, "Acc": 0.667, "F1": 0.739},
    "EfficientNet-B0": {"AUC": 0.849, "Sens": 0.368, "Spec": 1.000, "Acc": 0.500, "F1": 0.513},
}

# --- Hiperparametre Evrimi (Git Gecmisi) ---
# Phase 0: Ilk commit (156c925)
phase0 = {
    "UNet++":          None,  # config.py'de DenseNet varsayildi
    "DenseNet-121":    {"LR": "1e-4", "WD": "1e-4", "DO": "0.40", "BS": "4", "EP": "100", "ES": "20", "WU": "-", "LS": "-", "FG": "1.5", "FA": "0.50", "MGN": "1.0", "AUG": "Yok"},
    "EfficientNet-B0": None,
}

# Phase 1: Threshold iyilestirme sonrasi (4df2c9c)
phase1_changes = "LR: 5e-6 -> 1e-3, FG: 1.5 -> 2.0, Warmup eklendi (5 epoch)"

# Phase 2: Label Smoothing eklenmesi (2c69fa7)
phase2_changes = "Label Smoothing = 0.1 eklendi, Dropout: 0.4 -> 0.5"

# Phase 3: Model bazli ayarlar (cc4a3f1 - 5257bc2)
# Bu asama config.py'de yorum satiri ile 3 model profili olusturuldu

# Phase 4: Augmentasyon + Multi-Run Protocol (f86b4bf)
# Tamamen yeniden yapilandirildi

# =====================================================================
# GRAFIK 1: AUGMENTASYONSUZ Karsilastirma
# =====================================================================
def create_chart_no_aug():
    metrics = ["AUC", "Sens", "Spec", "F1"]
    labels = ["AUC-ROC", "Duyarlilik\n(Sensitivity)", "Ozgulluk\n(Specificity)", "F1-Score"]
    x = np.arange(len(metrics))
    width = 0.22
    
    fig, ax = plt.subplots(figsize=(12, 6))
    colors = ['#2196F3', '#FF9800', '#4CAF50']
    
    for i, m in enumerate(models):
        vals = [aug_yok[m][k] for k in metrics]
        bars = ax.bar(x + i*width, vals, width, label=m, color=colors[i], edgecolor='white', linewidth=0.5)
        for bar, val in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.02,
                    f'{val:.3f}', ha='center', va='bottom', fontsize=9, fontweight='bold')
    
    ax.set_xlabel('Metrik', fontsize=12, fontweight='bold')
    ax.set_ylabel('Deger', fontsize=12, fontweight='bold')
    ax.set_title('Augmentasyonsuz (Eski) Dis Test Sonuclari', fontsize=14, fontweight='bold', pad=15)
    ax.set_xticks(x + width)
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_ylim(0, 1.15)
    ax.legend(loc='upper left', fontsize=10, framealpha=0.9)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    # Onemli not
    ax.annotate('Model Cokmesi!\n(Tum hastalara\nPozitif dedi)', 
                xy=(2 + 2*width, 0.000), xytext=(2 + 2*width + 0.3, 0.25),
                fontsize=8, color='red', fontweight='bold',
                arrowprops=dict(arrowstyle='->', color='red', lw=1.5),
                ha='center')
    
    plt.tight_layout()
    path = os.path.join(OUT_DIR, "grafik_augmentasyonsuz_karsilastirma.png")
    plt.savefig(path, dpi=200, bbox_inches='tight')
    plt.close()
    return path

# =====================================================================
# GRAFIK 2: AUGMENTASYONLU Karsilastirma
# =====================================================================
def create_chart_with_aug():
    metrics = ["AUC", "Sens", "Spec", "F1"]
    labels = ["AUC-ROC", "Duyarlilik\n(Sensitivity)", "Ozgulluk\n(Specificity)", "F1-Score"]
    x = np.arange(len(metrics))
    width = 0.22
    
    fig, ax = plt.subplots(figsize=(12, 6))
    colors = ['#2196F3', '#FF9800', '#4CAF50']
    
    for i, m in enumerate(models):
        vals = [aug_var[m][k] for k in metrics]
        bars = ax.bar(x + i*width, vals, width, label=m, color=colors[i], edgecolor='white', linewidth=0.5)
        for bar, val in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.02,
                    f'{val:.3f}', ha='center', va='bottom', fontsize=9, fontweight='bold')
    
    ax.set_xlabel('Metrik', fontsize=12, fontweight='bold')
    ax.set_ylabel('Deger', fontsize=12, fontweight='bold')
    ax.set_title('Augmentasyonlu (3-Run Ortalama) Dis Test Sonuclari', fontsize=14, fontweight='bold', pad=15)
    ax.set_xticks(x + width)
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_ylim(0, 1.15)
    ax.legend(loc='upper left', fontsize=10, framealpha=0.9)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    path = os.path.join(OUT_DIR, "grafik_augmentasyonlu_karsilastirma.png")
    plt.savefig(path, dpi=200, bbox_inches='tight')
    plt.close()
    return path

# =====================================================================
# GRAFIK 3: AUGMENTASYONLU vs AUGMENTASYONSUZ (Birlestirme)
# =====================================================================
def create_chart_combined():
    metrics = ["AUC", "Sens", "Spec"]
    labels_short = ["AUC", "SENS", "SPEC"]
    
    fig, axes = plt.subplots(1, 3, figsize=(16, 6), sharey=True)
    colors_noaug = ['#90CAF9', '#FFCC80', '#A5D6A7']  # Acik tonlar
    colors_aug =   ['#1565C0', '#E65100', '#2E7D32']    # Koyu tonlar
    
    for idx, m in enumerate(models):
        ax = axes[idx]
        x = np.arange(len(metrics))
        width = 0.30
        
        vals_noaug = [aug_yok[m][k] for k in metrics]
        vals_aug   = [aug_var[m][k] for k in metrics]
        
        bars1 = ax.bar(x - width/2, vals_noaug, width, label='Augmentasyonsuz', 
                       color=colors_noaug[idx], edgecolor='gray', linewidth=0.5)
        bars2 = ax.bar(x + width/2, vals_aug, width, label='Augmentasyonlu', 
                       color=colors_aug[idx], edgecolor='gray', linewidth=0.5)
        
        for bar, val in zip(bars1, vals_noaug):
            ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.02,
                    f'{val:.3f}', ha='center', va='bottom', fontsize=8, fontweight='bold')
        for bar, val in zip(bars2, vals_aug):
            ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.02,
                    f'{val:.3f}', ha='center', va='bottom', fontsize=8, fontweight='bold')
        
        ax.set_title(m, fontsize=13, fontweight='bold', pad=10)
        ax.set_xticks(x)
        ax.set_xticklabels(labels_short, fontsize=10)
        ax.set_ylim(0, 1.2)
        ax.grid(axis='y', alpha=0.3, linestyle='--')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        if idx == 0:
            ax.set_ylabel('Deger', fontsize=12, fontweight='bold')
        ax.legend(fontsize=8, loc='upper right')
    
    fig.suptitle('Augmentasyonsuz vs Augmentasyonlu: Model Bazli Karsilastirma', 
                 fontsize=15, fontweight='bold', y=1.02)
    plt.tight_layout()
    path = os.path.join(OUT_DIR, "grafik_augmentasyonlu_vs_augmentasyonsuz.png")
    plt.savefig(path, dpi=200, bbox_inches='tight')
    plt.close()
    return path

# =====================================================================
# GRAFIK URETIMI
# =====================================================================
print("Grafikler olusturuluyor...")
chart1 = create_chart_no_aug()
chart2 = create_chart_with_aug()
chart3 = create_chart_combined()
print(f"  -> {chart1}")
print(f"  -> {chart2}")
print(f"  -> {chart3}")

# =====================================================================
# DOCX RAPOR OLUSTURMA
# =====================================================================
print("\nDOCX rapor hazirlaniyor...")
doc = Document()

# --- KAPAK ---
doc.add_heading("3D CT Siniflandirma Projesi\nKapsamli Hiperparametre ve Model Analiz Raporu", 0)
doc.add_paragraph("Bu rapor, projenin basindan sonuna kadar uygulanan tum hiperparametre ayarlarini, augmentasyonlu ve augmentasyonsuz egitim sonuclarini, grafik karsilastirmalarini ve her modelin bireysel analizini icerir. Tum veriler projenin Git gecmisinden (commit bazli) cikarilmis ve dogrulanmistir.")

# =====================================================================
# BOLUM 1: HIPERPARAMETRE EVRIMI
# =====================================================================
doc.add_heading("1. Hiperparametre Evrimi (Tum Degisiklikler - Git Gecmisi)", level=1)

doc.add_heading("Faz 0: Ilk Kurulum (commit: 156c925)", level=2)
doc.add_paragraph("Projenin ilk commitinde sadece DenseNet-121 modeli tanimlanmisti. Augmentasyon yoktu, Label Smoothing yoktu. Temel parametreler:")

t0 = doc.add_table(rows=2, cols=8)
t0.style = "Table Grid"
h0 = t0.rows[0].cells
h0[0].text, h0[1].text, h0[2].text, h0[3].text = "Model", "LR", "Weight Decay", "Dropout"
h0[4].text, h0[5].text, h0[6].text, h0[7].text = "Focal Gamma", "Focal Alpha", "Epochs", "Augmentation"
for c in h0: c.paragraphs[0].runs[0].bold = True
r0 = t0.rows[1].cells
r0[0].text, r0[1].text, r0[2].text, r0[3].text = "DenseNet-121", "1e-4", "1e-4", "0.40"
r0[4].text, r0[5].text, r0[6].text, r0[7].text = "1.5", "0.50", "100", "Yok"

doc.add_paragraph("")
doc.add_heading("Faz 1: Threshold Iyilestirme (commit: 4df2c9c)", level=2)
doc.add_paragraph("Model asiri yavas ogrendigi icin LR 5e-6'dan 1e-3'e yukseldi. Focal gamma 1.5'ten 2.0'a arttirildi (zor orneklere daha cok odaklanma). Warmup epoch eklendi.")
t1 = doc.add_table(rows=2, cols=4)
t1.style = "Table Grid"
h1 = t1.rows[0].cells
h1[0].text, h1[1].text, h1[2].text, h1[3].text = "Degisiklik", "Onceki", "Sonraki", "Neden"
for c in h1: c.paragraphs[0].runs[0].bold = True
t1.rows[1].cells[0].text = "LR / FG / Warmup"
t1.rows[1].cells[1].text = "LR=5e-6, FG=1.5, WU=Yok"
t1.rows[1].cells[2].text = "LR=1e-3, FG=2.0, WU=5"
t1.rows[1].cells[3].text = "Model ogrenmeye baslasin diye"

doc.add_paragraph("")
doc.add_heading("Faz 2: Label Smoothing Eklenmesi (commit: 2c69fa7)", level=2)
doc.add_paragraph("Model asiri guvenerek (overconfident) %100 precision verirken sensitivity dusuk kaliyordu. Bu sorunu cozmek icin Label Smoothing = 0.1 eklendi ve Dropout 0.4'ten 0.5'e artirildi.")
t2 = doc.add_table(rows=2, cols=4)
t2.style = "Table Grid"
h2 = t2.rows[0].cells
h2[0].text, h2[1].text, h2[2].text, h2[3].text = "Degisiklik", "Onceki", "Sonraki", "Neden"
for c in h2: c.paragraphs[0].runs[0].bold = True
t2.rows[1].cells[0].text = "LS / Dropout"
t2.rows[1].cells[1].text = "LS=Yok, DO=0.40"
t2.rows[1].cells[2].text = "LS=0.10, DO=0.50"
t2.rows[1].cells[3].text = "Overconfidence onleme"

doc.add_paragraph("")
doc.add_heading("Faz 3: Model Bazli Profiller (commit: cc4a3f1 / 5257bc2 / 9e5950b)", level=2)
doc.add_paragraph("3 model icin ayri ayri hiperparametre profilleri olusturuldu. Config.py dosyasinda yorum satirlari ile model secimi yapildi. Bu asama augmentasyonsuz son haldir.")

t3 = doc.add_table(rows=4, cols=11)
t3.style = "Table Grid"
h3 = t3.rows[0].cells
headers = ["Model", "LR", "WD", "Dropout", "Batch", "Epochs", "Patience", "Warmup", "Label Sm.", "Focal G", "Grad Norm"]
for i, h in enumerate(headers): 
    h3[i].text = h
    h3[i].paragraphs[0].runs[0].bold = True

data3 = [
    ["UNet++",          "1e-3", "1e-4", "0.50", "4", "100", "40", "5",  "0.05", "2.0", "1.0"],
    ["DenseNet-121",    "5e-4", "5e-4", "0.40", "4", "150", "30", "5",  "0.10", "2.0", "1.0"],
    ["EfficientNet-B0", "3e-4", "1e-3", "0.30", "8", "150", "25", "10", "0.10", "2.0", "0.5"],
]
for i, row in enumerate(data3):
    for j, val in enumerate(row):
        t3.rows[i+1].cells[j].text = val

doc.add_paragraph("")
doc.add_heading("Faz 4: Augmentasyon + Multi-Run Protocol (commit: f86b4bf / 3382929)", level=2)
doc.add_paragraph("Bilimsel olarak Fair Comparison (Adil Kiyaslama) yapabilmek icin tum hiperparametreler standartlastirildi. "
                  "9 Asamali 3D Augmentasyon eklendi. 3 Bagimsiz Run (seed: 42, 123, 456) ile 5-Fold CV uygulanmaya baslandi. "
                  "Focal gamma 2.0'dan 1.0'a dusuruldu (augmentasyon overfitting'i zaten onledigi icin asiri odaklanmaya gerek kalmadi). "
                  "Label Smoothing tum modellerde 0.05'e sabitlendi.")

t4 = doc.add_table(rows=4, cols=11)
t4.style = "Table Grid"
h4 = t4.rows[0].cells
for i, h in enumerate(headers):
    h4[i].text = h
    h4[i].paragraphs[0].runs[0].bold = True

data4 = [
    ["UNet++",          "1e-3", "5e-3", "0.40", "2 (eff:16)", "100", "25", "5",  "0.05", "1.0", "1.0"],
    ["DenseNet-121",    "1e-4", "5e-3", "0.40", "2 (eff:16)", "100", "25", "10", "0.05", "1.0", "1.0"],
    ["EfficientNet-B0", "1e-4", "1e-2", "0.40", "2 (eff:16)", "100", "25", "10", "0.05", "1.0", "1.0"],
]
for i, row in enumerate(data4):
    for j, val in enumerate(row):
        t4.rows[i+1].cells[j].text = val

doc.add_paragraph("")

# =====================================================================
# BOLUM 2: SONUCLAR (AUGMENTASYONSUZ)
# =====================================================================
doc.add_heading("2. Augmentasyonsuz Dis Test Sonuclari", level=1)
doc.add_paragraph("Asagidaki tablo, augmentasyon yapilmadan egitilen modellerin Dis Test (External Test) kumesindeki performansini gostermektedir.")

ts1 = doc.add_table(rows=4, cols=6)
ts1.style = "Table Grid"
hs1 = ts1.rows[0].cells
hs1[0].text, hs1[1].text, hs1[2].text, hs1[3].text, hs1[4].text, hs1[5].text = \
    "Model", "AUC-ROC", "Duyarlilik", "Ozgulluk", "F1-Score", "Dogruluk"
for c in hs1: c.paragraphs[0].runs[0].bold = True

for i, m in enumerate(models):
    d = aug_yok[m]
    row = ts1.rows[i+1].cells
    row[0].text = m
    row[1].text = f"{d['AUC']:.3f}"
    row[2].text = f"{d['Sens']:.3f}"
    row[3].text = f"{d['Spec']:.3f}"
    row[4].text = f"{d['F1']:.3f}"
    row[5].text = f"{d['Acc']:.3f}"

doc.add_paragraph("")
doc.add_paragraph("").add_run("Grafik 1: Augmentasyonsuz Model Karsilastirmasi").bold = True
doc.add_picture(chart1, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# =====================================================================
# BOLUM 3: SONUCLAR (AUGMENTASYONLU)
# =====================================================================
doc.add_heading("3. Augmentasyonlu (3-Run Ortalama) Dis Test Sonuclari", level=1)
doc.add_paragraph("Asagidaki tablo, 9 Asamali 3D Augmentasyon ve Multi-Run Protocol ile egitilen modellerin 3 bagimsiz calismanin ortalamasidir.")

ts2 = doc.add_table(rows=4, cols=6)
ts2.style = "Table Grid"
hs2 = ts2.rows[0].cells
hs2[0].text, hs2[1].text, hs2[2].text, hs2[3].text, hs2[4].text, hs2[5].text = \
    "Model", "Ort. AUC", "Ort. Duyarlilik", "Ort. Ozgulluk", "Ort. F1", "Ort. Dogruluk"
for c in hs2: c.paragraphs[0].runs[0].bold = True

for i, m in enumerate(models):
    d = aug_var[m]
    row = ts2.rows[i+1].cells
    row[0].text = m
    row[1].text = f"{d['AUC']:.3f}"
    row[2].text = f"{d['Sens']:.3f}"
    row[3].text = f"{d['Spec']:.3f}"
    row[4].text = f"{d['F1']:.3f}"
    row[5].text = f"{d['Acc']:.3f}"

doc.add_paragraph("")
doc.add_paragraph("").add_run("Grafik 2: Augmentasyonlu Model Karsilastirmasi").bold = True
doc.add_picture(chart2, width=Inches(5.5))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# =====================================================================
# BOLUM 4: AUGMENTASYONLU vs AUGMENTASYONSUZ
# =====================================================================
doc.add_heading("4. Augmentasyonlu vs Augmentasyonsuz: Tarafsiz Karsilastirma", level=1)

doc.add_paragraph("").add_run("Grafik 3: Tum Modeller - Augmentasyonlu vs Augmentasyonsuz").bold = True
doc.add_picture(chart3, width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# Karsilastirma tablosu
tc = doc.add_table(rows=4, cols=7)
tc.style = "Table Grid"
hc = tc.rows[0].cells
hc[0].text = "Model"
hc[1].text = "AUC (Eski)"
hc[2].text = "AUC (Yeni)"
hc[3].text = "Sens (Eski)"
hc[4].text = "Sens (Yeni)"
hc[5].text = "Spec (Eski)"
hc[6].text = "Spec (Yeni)"
for c in hc: c.paragraphs[0].runs[0].bold = True

for i, m in enumerate(models):
    r = tc.rows[i+1].cells
    r[0].text = m
    r[1].text = f"{aug_yok[m]['AUC']:.3f}"
    r[2].text = f"{aug_var[m]['AUC']:.3f}"
    r[3].text = f"{aug_yok[m]['Sens']:.3f}"
    r[4].text = f"{aug_var[m]['Sens']:.3f}"
    r[5].text = f"{aug_yok[m]['Spec']:.3f}"
    r[6].text = f"{aug_var[m]['Spec']:.3f}"

doc.add_paragraph("")

p_comp = doc.add_paragraph()
p_comp.add_run("Tarafsiz Degerlendirme:\n").bold = True
p_comp.add_run(
    "Augmentasyonlu ve augmentasyonsuz yaklasimlar farkli ticari dengeleri (trade-off) hedefler. "
    "Augmentasyonsuz modeller genel olarak daha yuksek Duyarlilik (Sensitivity) uretmistir "
    "cunku modeller egitim verisini daha cok ezberlemis ve bu da test setinde dogru teshisleri artirmistir. "
    "Ancak bu ezberleme, Ozgulluk (Specificity) tarafinda buyuk bir bedelle gelmistir: "
    "EfficientNet-B0 modeli %0 Ozgulluk ile tamamen cokmustur.\n\n"
    "Augmentasyonlu modeller ise cok daha dengeli ve guvenilir sonuclar vermistir. "
    "Ozgulluk oranlari belirgin sekilde yukselmis (UNet++ %80 -> %93.3, EfficientNet-B0 %0 -> %100) "
    "ve Yanlis Pozitif (False Positive) sayisi neredeyse sifira inmistir. "
    "Tibbi yapay zekada saglikli bir hastaya yanlis kanser/tumor teshisi koymak (False Positive) "
    "en istenmeyen durumlardan biridir. Bu acilardan augmentasyon buyuk bir basari getirmistir.\n\n"
    "Dezavantaji ise Duyarlilik (Sensitivity) degerlerinin dusuk olmasi: "
    "modeller artik asiri temkinli davranmakta ve bazi gercek pozitif hastalari kacirabilecegidir. "
    "Bu durum, klinik uygulamada ikinci bir dogrulama testine ihtiyac dogurur.\n\n"
    "Sonuc olarak: Augmentasyon, ozellikle kucuk tibbi veri setlerinde overfitting'i onlemek ve "
    "modellerin gercek dunyada guvenilir calismasi icin kesinlikle gereklidir. "
    "Ancak ideal denge, augmentasyon ile birlikte model mimarisinin dogru secilmesine baglidir."
)

# =====================================================================
# BOLUM 5: MODEL BAZLI DETAYLI ANALIZ
# =====================================================================
doc.add_heading("5. Model Bazli Detayli Analiz", level=1)

# --- UNet++ ---
doc.add_heading("5.1 UNet++ (3D Custom CNN - ~890K Parametre)", level=2)
p_u = doc.add_paragraph()
p_u.add_run("Mimari Ozellikleri: ").bold = True
p_u.add_run("UNet++ mimarisi, projede sifirdan tasarlanmis kompakt bir 3D Convolutional Neural Network'tur. "
            "Yaklaik 890 bin parametreye sahiptir. Encoder-Decoder yapisinda derin yuvalenmis (nested) atlama baglantilari (skip connections) kullanir. "
            "Bu sayede hem ust duzey (sematic) hem de alt duzey (detail) ozellikleri birlikte degerlendirerek karar verir.\n\n")

p_u.add_run("Agirlik Ayarlari ve Degisiklikler:\n").bold = True
p_u.add_run("- Ilk Asama: UNet++ projede ilk kullanilan model degildir; baslangicta DenseNet-121 varsayildi. "
            "Sonrasinda model cokmesi (Mode Collapse) yasamasi uzerine UNet++'a gecildi (commit 9e5950b).\n"
            "- Augmentasyonsuz Donem: LR=1e-3, WD=1e-4, Dropout=0.50, Label Smoothing=0.05, Focal Gamma=2.0 ile egitildi. "
            "En iyi AUC (0.863) ve en dengeli Sensitivity/Specificity orani (0.789/0.800) bu donemde elde edildi.\n"
            "- Augmentasyonlu Donem: WD 1e-4'ten 5e-3'e artirildi (daha guclu L2 regülarizasyonu). Dropout 0.50'den 0.40'a indirildi "
            "(augmentasyon zaten kendi basina regularizerdir). Focal gamma 2.0'dan 1.0'a dusuruldu.\n\n")

p_u.add_run("Sonuc Degerlendirmesi:\n").bold = True
p_u.add_run("UNet++ hem augmentasyonsuz hem augmentasyonlu donemlerde en yuksek AUC skorunu (0.863 ve 0.853) ureten modeldir. "
            "Diger modellerin aksine hicbir donemde cokmemis, istikrarli sonuclar vermistir. "
            "Kucuk parametre sayisi sayesinde sinirli tibbi veri setlerinde asiri ogrenmeden kacinabilmektedir.")

# --- DenseNet-121 ---
doc.add_heading("5.2 DenseNet-121 (Transfer Learning - ~11.7M Parametre)", level=2)
p_d = doc.add_paragraph()
p_d.add_run("Mimari Ozellikleri: ").bold = True
p_d.add_run("DenseNet-121, her katmanin onceki tum katmanlara baglandigi Dense Block yapisiyla taninir. "
            "Yaklaik 11.7 milyon parametreye sahiptir. ImageNet uzerinde onceden egitilmis (pre-trained) agirliklar kullanilarak "
            "transfer learning ile 3D CT verilerine ince ayar (fine-tune) yapilmistir. "
            "2D'den 3D'ye gecis icin ilk konvolusyon katmani uyarlanmistir.\n\n")

p_d.add_run("Agirlik Ayarlari ve Degisiklikler:\n").bold = True
p_d.add_run("- Ilk Asama (Faz 0): Projenin varsayilan modeli olarak LR=1e-4, WD=1e-4, Focal Gamma=1.5 ile basladi.\n"
            "- Faz 1-2: LR 1e-3'e cikti, Focal Gamma 2.0'a artti, Label Smoothing=0.10 eklendi.\n"
            "- Augmentasyonsuz Son Hal: LR=5e-4, WD=5e-4, Dropout=0.40, LS=0.10, FG=2.0, Epochs=150 olarak ayarlandi.\n"
            "- Augmentasyonlu Donem: LR=1e-4'e dusuruld, WD=5e-3'e artirildi, LS=0.05 ve FG=1.0 olarak standartlastirildi.\n\n")

p_d.add_run("Sonuc Degerlendirmesi:\n").bold = True
p_d.add_run("DenseNet-121, augmentasyonsuz donemde AUC=0.842 ile UNet++'in gerisinde kalmistir. "
            "Duyarlilik (%84.2) yuksek gorunse de, Ozgulluk (%60.0) onemli olcude dusuktur - bu da 5 saglikli hastanin 2'sine yanlis kanser teshisi konuldugu anlamina gelir. "
            "Augmentasyonlu donemde AUC=0.849 ile UNet++'a (0.853) yakinlasmistir ancak parametre sayisi (11.7M vs 890K) goz onune alindiginda "
            "bu kucuk fark kayda deger degildir ve UNet++ cok daha verimlidir.")

# --- EfficientNet-B0 ---
doc.add_heading("5.3 EfficientNet-B0 (Transfer Learning - ~4.7M Parametre)", level=2)
p_e = doc.add_paragraph()
p_e.add_run("Mimari Ozellikleri: ").bold = True
p_e.add_run("EfficientNet-B0, Google'in Neural Architecture Search (NAS) ile otomatik olarak optimize ettigi bir mimaridir. "
            "Compound Scaling yontemiyle derinlik, genislik ve cozunurlugu birlikte olceklendirir. "
            "Yaklaik 4.7 milyon parametreye sahiptir. ImageNet onceden egitilmis agirliklar ile transfer learning uygulanmistir.\n\n")

p_e.add_run("Agirlik Ayarlari ve Degisiklikler:\n").bold = True
p_e.add_run("- Augmentasyonsuz Donem: LR=3e-4, WD=1e-3, Dropout=0.30, LS=0.10, FG=2.0, Max Grad Norm=0.5 ile egitildi. "
            "Farkli bir donemde loss fonksiyonu CrossEntropy'ye cevrilerek (commit cc4a3f1) LR=1e-4, WD=1e-4 ile de denenmistir.\n"
            "- Augmentasyonlu Donem: LR=1e-4, WD=1e-2 (en agir regülarizasyon), Dropout=0.40, LS=0.05, FG=1.0 olarak ayarlandi.\n\n")

p_e.add_run("Sonuc Degerlendirmesi:\n").bold = True
p_e.add_run("EfficientNet-B0, projenin en sorunlu modeli olmustur. Augmentasyonsuz donemde tamamen cokmustur (Mode Collapse): "
            "Duyarlilik %100, Ozgulluk %0. Bu, modelin egitim verisindeki sinif dengesizligini (19 Mukozlu vs 5 Apandisit) tamamen ezberleyerek "
            "tum hastalara 'Mukozlu' demesi anlamina gelir. Klinik olarak kullanisizdir.\n\n"
            "Augmentasyonlu donemde cokme problemi cozulmustur (Ozgulluk %100) ancak bu sefer Duyarlilik %36.8'e dusmustur. "
            "Model asiri temkinli hale gelerek gercek pozitif hastalari kacirmaktadir. "
            "Bu durum, EfficientNet-B0'in bu kucuk veri seti icin uygun bir mimari olmadigini gostermektedir.")

# =====================================================================
# BOLUM 6: GENEL SONUC
# =====================================================================
doc.add_heading("6. Genel Sonuc ve Oneriler", level=1)
p_sonuc = doc.add_paragraph()
p_sonuc.add_run("1. Augmentasyon Kesinlikle Gereklidir: ").bold = True
p_sonuc.add_run("Kucuk tibbi veri setlerinde (140 hasta) augmentasyonsuz egitim, modellerin veriyi ezberlemesine ve cokmes"
                "ine yol acmaktadir. EfficientNet-B0'in augmentasyonsuz %0 Ozgulluk uretmesi bunun en net kanitidir.\n\n")

p_sonuc.add_run("2. En Basarili Model: UNet++: ").bold = True
p_sonuc.add_run("Her iki donemde de (augmentasyonlu ve augmentasyonsuz) en yuksek AUC'yi ureten, en dengeli Duyarlilik/Ozgulluk oranini saglayan "
                "ve hicbir zaman cokmeyen tek model UNet++ olmustur. Kompakt yapisi (890K parametre) kucuk veri setlerine cok uygundur.\n\n")

p_sonuc.add_run("3. DenseNet-121: Idare Eder Ancak Hantal: ").bold = True
p_sonuc.add_run("UNet++'a yakin sonuclar uretmesine ragmen 13 kat daha fazla parametreye sahiptir (11.7M vs 890K). "
                "Egitim suresi, hafiza tuketimi ve kompleksite acisindan avantaji yoktur.\n\n")

p_sonuc.add_run("4. EfficientNet-B0: Bu Veri Seti Icin Uygun Degil: ").bold = True
p_sonuc.add_run("Her iki donemde de problemlidir. Ya tamamen cokmustur ya da asiri temkinli hale gelmistir. "
                "Compound Scaling yaklasimi buyuk veri setlerinde basariliyken, sinirli tibbi verilerde dezavantaj yaratmaktadir.")

# KAYDET
out_path = os.path.join(OUT_DIR, "Kapsamli_Hiperparametre_ve_Model_Analiz_Raporu.docx")
doc.save(out_path)
print(f"\nRapor kaydedildi: {out_path}")
print("Tamamlandi!")
