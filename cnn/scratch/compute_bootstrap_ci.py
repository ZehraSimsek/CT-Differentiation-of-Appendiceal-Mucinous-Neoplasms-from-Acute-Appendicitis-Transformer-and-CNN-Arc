#!/usr/bin/env python3
"""
Bootstrap %95 Guven Araliklari (CI) Hesaplama
==============================================
Her model icin 3 bagimsiz run'in hasta bazli tahminlerini
birlestirerek (pooled) 2000 bootstrap orneklemesiyle
AUC, Sensitivity, Specificity, F1, Accuracy icin %95 CI uretir.

Ayrica run-level variabilite (3 run ortalamalari +/- std) de hesaplanir.
"""

import os
import numpy as np
import pandas as pd
from sklearn.metrics import roc_auc_score, f1_score, accuracy_score

np.random.seed(42)

# =====================================================================
# YAPILANDIRMA
# =====================================================================
N_BOOTSTRAP = 2000       # MULTI_RUN_PROTOCOL.md'de taahhut edilen sayi
CI_ALPHA     = 0.05      # %95 CI
OUT_DIR      = "paper_tables_and_curves"

PATIENT_CSV  = os.path.join(OUT_DIR, "tum_test_ve_fold_verileri_csv",
                            "master_external_test_patient_predictions.csv")
RUN_CSV      = os.path.join(OUT_DIR, "tum_test_ve_fold_verileri_csv",
                            "master_run_level_metrics.csv")

# =====================================================================
# VERI YUKLEME
# =====================================================================
df_pat = pd.read_csv(PATIENT_CSV)
df_run = pd.read_csv(RUN_CSV)

models = [
    ("unet_plusplus",    "UNet++"),
    ("densenet121",     "DenseNet-121"),
    ("efficientnet_b0", "EfficientNet-B0"),
]

# =====================================================================
# YARDIMCI FONKSIYONLAR
# =====================================================================
def compute_metrics(y_true, y_prob, threshold=0.5):
    """Tek bir metrik seti hesapla."""
    y_pred = (y_prob >= threshold).astype(int)
    
    tp = np.sum((y_pred == 1) & (y_true == 1))
    tn = np.sum((y_pred == 0) & (y_true == 0))
    fp = np.sum((y_pred == 1) & (y_true == 0))
    fn = np.sum((y_pred == 0) & (y_true == 1))
    
    sens = tp / (tp + fn) if (tp + fn) > 0 else 0.0
    spec = tn / (tn + fp) if (tn + fp) > 0 else 0.0
    prec = tp / (tp + fp) if (tp + fp) > 0 else 0.0
    f1   = 2*prec*sens / (prec + sens) if (prec + sens) > 0 else 0.0
    acc  = (tp + tn) / (tp + tn + fp + fn)
    
    try:
        auc = roc_auc_score(y_true, y_prob)
    except ValueError:
        auc = np.nan
    
    return {
        "AUC": auc,
        "Sensitivity": sens,
        "Specificity": spec,
        "Precision": prec,
        "F1": f1,
        "Accuracy": acc,
        "TP": int(tp), "TN": int(tn), "FP": int(fp), "FN": int(fn),
    }


def bootstrap_ci(y_true, y_prob, n_bootstrap=2000, ci_alpha=0.05, threshold=0.5):
    """Stratified bootstrap ile %95 CI hesapla."""
    n = len(y_true)
    metric_names = ["AUC", "Sensitivity", "Specificity", "Precision", "F1", "Accuracy"]
    boot_results = {m: [] for m in metric_names}
    
    for _ in range(n_bootstrap):
        # Stratified resampling: sinif oranlarini koru
        idx_pos = np.where(y_true == 1)[0]
        idx_neg = np.where(y_true == 0)[0]
        
        boot_pos = np.random.choice(idx_pos, size=len(idx_pos), replace=True)
        boot_neg = np.random.choice(idx_neg, size=len(idx_neg), replace=True)
        boot_idx = np.concatenate([boot_pos, boot_neg])
        
        y_t = y_true[boot_idx]
        y_p = y_prob[boot_idx]
        
        metrics = compute_metrics(y_t, y_p, threshold)
        for m in metric_names:
            boot_results[m].append(metrics[m])
    
    ci = {}
    for m in metric_names:
        arr = np.array(boot_results[m])
        arr = arr[~np.isnan(arr)]
        if len(arr) > 0:
            lower = np.percentile(arr, 100 * ci_alpha / 2)
            upper = np.percentile(arr, 100 * (1 - ci_alpha / 2))
            mean  = np.mean(arr)
            ci[m] = {"mean": mean, "lower": lower, "upper": upper}
        else:
            ci[m] = {"mean": np.nan, "lower": np.nan, "upper": np.nan}
    
    return ci


# =====================================================================
# ANA HESAPLAMA
# =====================================================================
results = []

for model_id, model_name in models:
    print(f"\n{'='*60}")
    print(f"  {model_name} - Bootstrap %95 CI (n={N_BOOTSTRAP})")
    print(f"{'='*60}")
    
    # --- Yontem 1: Pooled Bootstrap (3 run x 24 hasta = 72 gozlem) ---
    sub = df_pat[df_pat["model_id"] == model_id]
    y_true_pooled = sub["true_label"].values
    y_prob_pooled = sub["ensemble_prob"].values
    
    # Temel metrikler (pooled)
    base_metrics = compute_metrics(y_true_pooled, y_prob_pooled)
    print(f"\n  Pooled Base Metrics (n={len(y_true_pooled)}):")
    for k, v in base_metrics.items():
        if k not in ("TP","TN","FP","FN"):
            print(f"    {k}: {v:.4f}")
    
    # Bootstrap CI (pooled)
    ci_pooled = bootstrap_ci(y_true_pooled, y_prob_pooled, 
                             n_bootstrap=N_BOOTSTRAP, ci_alpha=CI_ALPHA)
    
    print(f"\n  Bootstrap %95 CI (Pooled, n_boot={N_BOOTSTRAP}):")
    for metric_name, vals in ci_pooled.items():
        print(f"    {metric_name}: {vals['mean']:.4f} [{vals['lower']:.4f} - {vals['upper']:.4f}]")
        results.append({
            "Model": model_name,
            "Method": "Pooled Bootstrap",
            "Metric": metric_name,
            "Point_Estimate": base_metrics[metric_name],
            "Bootstrap_Mean": vals["mean"],
            "CI_Lower_2.5%": vals["lower"],
            "CI_Upper_97.5%": vals["upper"],
        })
    
    # --- Yontem 2: Run-Level Variabilite (3 run ortalamalari) ---
    sub_runs = df_run[df_run["model_id"] == model_id]
    
    run_metrics = {
        "AUC":         sub_runs["ext_auc"].values,
        "Sensitivity": sub_runs["ext_sens_05"].values,
        "Specificity": sub_runs["ext_spec_05"].values,
        "F1":          sub_runs["ext_f1_05"].values,
        "Accuracy":    sub_runs["ext_acc_05"].values,
    }
    
    print(f"\n  Run-Level Variabilite (3 Run):")
    for metric_name, vals in run_metrics.items():
        mean_val = np.mean(vals)
        std_val  = np.std(vals, ddof=1) if len(vals) > 1 else 0.0
        print(f"    {metric_name}: {mean_val:.4f} +/- {std_val:.4f}  (runs: {[f'{v:.4f}' for v in vals]})")
        results.append({
            "Model": model_name,
            "Method": "Run-Level (mean +/- SD)",
            "Metric": metric_name,
            "Point_Estimate": mean_val,
            "Bootstrap_Mean": mean_val,
            "CI_Lower_2.5%": mean_val - 1.96 * std_val,
            "CI_Upper_97.5%": mean_val + 1.96 * std_val,
        })
    
    # --- Yontem 3: Per-Run Bootstrap (her run icin ayri 24 hastalik CI) ---
    for run_idx in sorted(sub["run_idx"].unique()):
        run_sub = sub[sub["run_idx"] == run_idx]
        y_true_run = run_sub["true_label"].values
        y_prob_run = run_sub["ensemble_prob"].values
        
        ci_run = bootstrap_ci(y_true_run, y_prob_run,
                              n_bootstrap=N_BOOTSTRAP, ci_alpha=CI_ALPHA)
        
        print(f"\n  Run {run_idx} Bootstrap %95 CI (n={len(y_true_run)}):")
        for metric_name, vals in ci_run.items():
            print(f"    {metric_name}: {vals['mean']:.4f} [{vals['lower']:.4f} - {vals['upper']:.4f}]")
            results.append({
                "Model": model_name,
                "Method": f"Run {run_idx} Bootstrap",
                "Metric": metric_name,
                "Point_Estimate": compute_metrics(y_true_run, y_prob_run)[metric_name],
                "Bootstrap_Mean": vals["mean"],
                "CI_Lower_2.5%": vals["lower"],
                "CI_Upper_97.5%": vals["upper"],
            })

# =====================================================================
# CSV KAYDETME
# =====================================================================
df_results = pd.DataFrame(results)
csv_path = os.path.join(OUT_DIR, "tum_test_ve_fold_verileri_csv", "bootstrap_95_ci_results.csv")
df_results.to_csv(csv_path, index=False)
print(f"\n\nCSV kaydedildi: {csv_path}")

# =====================================================================
# DOCX TABLO OLUSTURMA
# =====================================================================
from docx import Document
from docx.shared import Pt, Inches

doc = Document()
doc.add_heading("Bootstrap %95 Guven Araliklari (CI) Raporu", 0)
doc.add_paragraph(f"Bu rapor, 3 bagimsiz calismanin (Run 1-3) dis test seti (External Test, n=24 hasta) "
                  f"sonuclari uzerinden {N_BOOTSTRAP} stratified bootstrap orneklemesiyle hesaplanmis "
                  f"%95 guven araliklarini (CI) icerir.")

# --- Ana Tablo: Pooled Bootstrap ---
doc.add_heading("Tablo 1: Pooled Bootstrap %95 CI (3 Run x 24 Hasta = 72 Gozlem)", level=2)
doc.add_paragraph("Tum runlarin hasta bazli tahminleri birlestirilerek (pooled) bootstrap CI hesaplanmistir. "
                  "Bu yontem, modelin genel performansinin guven araligini gosterir.")

metric_names = ["AUC", "Sensitivity", "Specificity", "Precision", "F1", "Accuracy"]
t_main = doc.add_table(rows=1 + len(models), cols=1 + len(metric_names))
t_main.style = "Table Grid"

# Header
h = t_main.rows[0].cells
h[0].text = "Model"
for i, mn in enumerate(metric_names):
    h[i+1].text = mn
for c in h:
    c.paragraphs[0].runs[0].bold = True

# Data
for mi, (model_id, model_name) in enumerate(models):
    row = t_main.rows[mi+1].cells
    row[0].text = model_name
    
    for ni, mn in enumerate(metric_names):
        match = df_results[(df_results["Model"]==model_name) & 
                           (df_results["Method"]=="Pooled Bootstrap") &
                           (df_results["Metric"]==mn)]
        if not match.empty:
            r = match.iloc[0]
            row[ni+1].text = f"{r['Point_Estimate']:.3f}\n[{r['CI_Lower_2.5%']:.3f}-{r['CI_Upper_97.5%']:.3f}]"

doc.add_paragraph("")

# --- Per-Run Tablo ---
doc.add_heading("Tablo 2: Run Bazli Bootstrap %95 CI (Her Run n=24 Hasta)", level=2)
doc.add_paragraph("Her run icin ayri ayri 24 hasta uzerinden hesaplanmis bootstrap CI degerleri.")

for model_id, model_name in models:
    doc.add_heading(f"{model_name}", level=3)
    
    t_run = doc.add_table(rows=4, cols=1 + len(metric_names))
    t_run.style = "Table Grid"
    h_run = t_run.rows[0].cells
    h_run[0].text = "Run"
    for i, mn in enumerate(metric_names):
        h_run[i+1].text = mn
    for c in h_run:
        c.paragraphs[0].runs[0].bold = True
    
    for ri, run_idx in enumerate([1, 2, 3]):
        row = t_run.rows[ri+1].cells
        row[0].text = f"Run {run_idx}"
        
        for ni, mn in enumerate(metric_names):
            match = df_results[(df_results["Model"]==model_name) & 
                               (df_results["Method"]==f"Run {run_idx} Bootstrap") &
                               (df_results["Metric"]==mn)]
            if not match.empty:
                r = match.iloc[0]
                row[ni+1].text = f"{r['Point_Estimate']:.3f}\n[{r['CI_Lower_2.5%']:.3f}-{r['CI_Upper_97.5%']:.3f}]"
    
    doc.add_paragraph("")

# --- Makale icin Formatlenmis Ozet Tablo ---
doc.add_heading("Tablo 3: Makale Tablosu Formati (Kopyala-Yapistir)", level=2)
doc.add_paragraph("Asagidaki tablo, dogrudan makaleye eklenebilecek formattadir. "
                  "Degerler: Ortalama [%95 CI Alt - Ust]")

t_paper = doc.add_table(rows=4, cols=5)
t_paper.style = "Table Grid"
hp = t_paper.rows[0].cells
hp[0].text = "Model"
hp[1].text = "AUC-ROC [%95 CI]"
hp[2].text = "Sensitivity [%95 CI]"
hp[3].text = "Specificity [%95 CI]"
hp[4].text = "F1-Score [%95 CI]"
for c in hp:
    c.paragraphs[0].runs[0].bold = True

for mi, (model_id, model_name) in enumerate(models):
    row = t_paper.rows[mi+1].cells
    row[0].text = model_name
    
    for ni, (mn, col_idx) in enumerate([("AUC",1), ("Sensitivity",2), ("Specificity",3), ("F1",4)]):
        match = df_results[(df_results["Model"]==model_name) & 
                           (df_results["Method"]=="Pooled Bootstrap") &
                           (df_results["Metric"]==mn)]
        if not match.empty:
            r = match.iloc[0]
            row[col_idx].text = f"{r['Point_Estimate']:.3f} [{r['CI_Lower_2.5%']:.3f}-{r['CI_Upper_97.5%']:.3f}]"

docx_path = os.path.join(OUT_DIR, "Bootstrap_95_CI_Raporu.docx")
doc.save(docx_path)
print(f"DOCX kaydedildi: {docx_path}")
print("\nTamamlandi!")
